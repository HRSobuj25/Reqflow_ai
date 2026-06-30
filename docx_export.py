"""
docx_export.py — ReqFlow AI · Professional DOCX Export Engine
==============================================================
Exports ERP requirement documents (BRD, SRS, Use Cases, User Stories,
Database Suggestions, KPIs, Workflow, Reports) as beautifully formatted
Microsoft Word files (.docx).

Usage
-----
    from docx_export import build_brd_docx, build_complete_suite_docx, DOCX_MIME
    import streamlit as st

    brd_bytes = build_brd_docx(
        project_name="Zenith Supply Logistics",
        industry="Supply Chain",
        content=active_project_data["brd"],
        timestamp="Jun 30, 14:22",
    )
    st.download_button("Export BRD as DOCX", data=brd_bytes,
                       file_name="Zenith_BRD.docx", mime=DOCX_MIME)

Public API
----------
    DOCX_MIME                    str    correct MIME type for .docx files
    build_brd_docx()             bytes  single BRD document
    build_section_docx()         bytes  any single ERP section
    build_complete_suite_docx()  bytes  all 8 ERP sections in one file
"""

from __future__ import annotations

import io
import re
from datetime import datetime
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Emu

# ---------------------------------------------------------------------------
# Public constants
# ---------------------------------------------------------------------------

DOCX_MIME = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)

# ReqFlow brand colours (purple / indigo theme)
_CLR_PURPLE   = RGBColor(0x8B, 0x5C, 0xF6)   # #8B5CF6 primary accent
_CLR_INDIGO   = RGBColor(0x63, 0x66, 0xF1)   # #6366F1 secondary accent
_CLR_DARK     = RGBColor(0x1E, 0x1B, 0x4B)   # #1E1B4B dark heading
_CLR_MUTED    = RGBColor(0x4F, 0x46, 0xE5)   # #4F46E5 subheading
_CLR_BODY     = RGBColor(0x1F, 0x29, 0x37)   # #1F2937 body text
_CLR_DIM      = RGBColor(0x6B, 0x72, 0x80)   # #6B7280 captions / muted
_CLR_BORDER   = RGBColor(0xE5, 0xE7, 0xEB)   # #E5E7EB table borders
_CLR_TH_BG   = "EEE9FF"                       # light purple table header bg
_CLR_TH_TEXT  = RGBColor(0x5B, 0x21, 0xB6)   # #5B21B6

# Section labels (icon + title) used in complete suite
_SECTION_META: dict[str, tuple[str, str]] = {
    "brd":            ("Business Requirements Document",        "BRD"),
    "srs":            ("Software Requirements Specification",   "SRS"),
    "use_cases":      ("Use Case Specifications",               "UC"),
    "user_stories":   ("User Stories & Acceptance Criteria",    "US"),
    "db_suggestions": ("Database Design Suggestions",           "DB"),
    "kpis":           ("Key Performance Indicators",            "KPI"),
    "workflow":       ("Workflow Specifications",               "WF"),
    "reports":        ("Report Specifications",                 "RPT"),
}


# ===========================================================================
#  Low-level element helpers
# ===========================================================================

def _set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def _set_cell_border(cell, color: str = "E5E7EB") -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _set_para_border_bottom(para, color: str = "8B5CF6") -> None:
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_page_break(doc: Document) -> None:
    from docx.enum.text import WD_BREAK
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(WD_BREAK.PAGE)


def _add_horizontal_rule(doc: Document, color: str = "8B5CF6") -> None:
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    para.paragraph_format.space_after = Pt(4)


# ===========================================================================
#  Markdown → DOCX converter
# ===========================================================================

_RE_H1     = re.compile(r"^#\s+(.+)$")
_RE_H2     = re.compile(r"^##\s+(.+)$")
_RE_H3     = re.compile(r"^###\s+(.+)$")
_RE_H4     = re.compile(r"^####\s+(.+)$")
_RE_BULLET = re.compile(r"^[-*\u2022]\s+(.+)$")
_RE_NUM    = re.compile(r"^\d+\.\s+(.+)$")
_RE_HR     = re.compile(r"^[-_*]{3,}$")
_RE_CODE   = re.compile(r"^```")
_RE_TROW   = re.compile(r"^\|(.+)\|$")
_RE_TSEP   = re.compile(r"^\|[\s\-:|]+\|$")


def _apply_inline(para, text: str) -> None:
    """Parses **bold** and `code` inline markers and adds styled runs."""
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        bold_m = re.fullmatch(r"\*\*([^*]+)\*\*", part)
        code_m = re.fullmatch(r"`([^`]+)`", part)
        if bold_m:
            run = para.add_run(bold_m.group(1))
            run.font.bold = True
            run.font.color.rgb = _CLR_DARK
        elif code_m:
            run = para.add_run(code_m.group(1))
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = _CLR_MUTED
        else:
            run = para.add_run(part)
            run.font.color.rgb = _CLR_BODY


def _flush_table(doc: Document, lines: list) -> None:
    parsed = []
    for line in lines:
        if _RE_TSEP.match(line.strip()):
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        parsed.append(cells)
    if not parsed:
        return
    cols = max(len(r) for r in parsed)
    tbl = doc.add_table(rows=len(parsed), cols=cols)
    tbl.style = "Table Grid"
    for r_idx, row in enumerate(parsed):
        for c_idx in range(cols):
            cell = tbl.cell(r_idx, c_idx)
            _set_cell_border(cell)
            text = row[c_idx] if c_idx < len(row) else ""
            p = cell.paragraphs[0]
            if r_idx == 0:
                _set_cell_bg(cell, _CLR_TH_BG)
                run = p.add_run(text)
                run.font.bold = True
                run.font.size = Pt(9.5)
                run.font.color.rgb = _CLR_TH_TEXT
            else:
                run = p.add_run(text)
                run.font.size = Pt(9.5)
                run.font.color.rgb = _CLR_BODY
    doc.add_paragraph()


def _render_markdown(doc: Document, markdown: str) -> None:
    """Converts Markdown text to styled DOCX paragraphs."""
    lines = (markdown or "").splitlines()
    in_code = False
    table_buf: list = []

    for raw in lines:
        line = raw.rstrip()

        # ── Code fence toggle
        if _RE_CODE.match(line):
            if in_code:
                in_code = False
                doc.add_paragraph()
            else:
                if table_buf:
                    _flush_table(doc, table_buf)
                    table_buf.clear()
                in_code = True
            continue

        if in_code:
            p = doc.add_paragraph()
            r = p.add_run(line)
            r.font.name = "Courier New"
            r.font.size = Pt(9)
            r.font.color.rgb = _CLR_MUTED
            p.paragraph_format.left_indent = Inches(0.35)
            p.paragraph_format.space_after = Pt(0)
            continue

        # ── Table rows
        if _RE_TROW.match(line):
            table_buf.append(line)
            continue
        else:
            if table_buf:
                _flush_table(doc, table_buf)
                table_buf.clear()

        # ── Horizontal rule
        if _RE_HR.match(line.strip()):
            _add_horizontal_rule(doc)
            continue

        # ── Headings
        m = _RE_H1.match(line)
        if m:
            p = doc.add_heading(level=1)
            r = p.add_run(m.group(1))
            r.font.color.rgb = _CLR_DARK
            r.font.size = Pt(16)
            r.font.bold = True
            continue

        m = _RE_H2.match(line)
        if m:
            p = doc.add_heading(level=2)
            r = p.add_run(m.group(1))
            r.font.color.rgb = _CLR_PURPLE
            r.font.size = Pt(13)
            r.font.bold = True
            continue

        m = _RE_H3.match(line)
        if m:
            p = doc.add_heading(level=3)
            r = p.add_run(m.group(1))
            r.font.color.rgb = _CLR_MUTED
            r.font.size = Pt(11.5)
            r.font.bold = True
            continue

        m = _RE_H4.match(line)
        if m:
            p = doc.add_heading(level=4)
            r = p.add_run(m.group(1))
            r.font.color.rgb = _CLR_BODY
            r.font.size = Pt(10.5)
            r.font.bold = True
            continue

        # ── Lists
        m = _RE_BULLET.match(line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            _apply_inline(p, m.group(1))
            p.paragraph_format.space_after = Pt(2)
            continue

        m = _RE_NUM.match(line)
        if m:
            p = doc.add_paragraph(style="List Number")
            _apply_inline(p, m.group(1))
            p.paragraph_format.space_after = Pt(2)
            continue

        # ── Blank line
        if not line.strip():
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            continue

        # ── Body paragraph
        p = doc.add_paragraph()
        _apply_inline(p, line)
        p.paragraph_format.space_after = Pt(4)

    if table_buf:
        _flush_table(doc, table_buf)


# ===========================================================================
#  Document skeleton builders
# ===========================================================================

def _create_base_doc() -> Document:
    """Creates a Document with A4 margins and Calibri default font."""
    doc = Document()
    sec = doc.sections[0]
    sec.page_height    = Emu(15120000)   # A4
    sec.page_width     = Emu(10692000)
    sec.left_margin    = Inches(1.1)
    sec.right_margin   = Inches(1.1)
    sec.top_margin     = Inches(1.0)
    sec.bottom_margin  = Inches(1.0)

    normal = doc.styles["Normal"]
    normal.font.name      = "Calibri"
    normal.font.size      = Pt(10.5)
    normal.font.color.rgb = _CLR_BODY
    return doc


def _add_cover_page(
    doc: Document,
    project_name: str,
    industry: str,
    document_type: str,
    timestamp: str,
) -> None:
    """Renders a professional company-branded cover page."""
    doc.add_paragraph()   # top padding

    # Brand
    b = doc.add_paragraph()
    b.alignment = WD_ALIGN_PARAGRAPH.CENTER
    br = b.add_run("ReqFlow AI  |  ERP Requirement Engine")
    br.font.name = "Calibri"
    br.font.size = Pt(10)
    br.font.bold = True
    br.font.color.rgb = _CLR_PURPLE
    b.paragraph_format.space_after = Pt(6)

    # Document type chip
    dt = doc.add_paragraph()
    dt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = dt.add_run(document_type.upper())
    dr.font.name = "Calibri"
    dr.font.size = Pt(9)
    dr.font.bold = True
    dr.font.color.rgb = _CLR_DIM
    dt.paragraph_format.space_after = Pt(18)

    # Big title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_border_bottom(t)
    tr = t.add_run(project_name)
    tr.font.name = "Calibri"
    tr.font.size = Pt(32)
    tr.font.bold = True
    tr.font.color.rgb = _CLR_DARK
    t.paragraph_format.space_after = Pt(20)

    # Industry
    ind = doc.add_paragraph()
    ind.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ir = ind.add_run(f"Industry Domain: {industry}")
    ir.font.name = "Calibri"
    ir.font.size = Pt(11)
    ir.font.italic = True
    ir.font.color.rgb = _CLR_MUTED
    ind.paragraph_format.space_after = Pt(6)

    # Timestamp
    ts = doc.add_paragraph()
    ts.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tsr = ts.add_run(f"Generated: {timestamp}")
    tsr.font.name = "Calibri"
    tsr.font.size = Pt(9.5)
    tsr.font.color.rgb = _CLR_DIM
    ts.paragraph_format.space_after = Pt(30)

    # Confidentiality
    cf = doc.add_paragraph()
    cf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cfr = cf.add_run(
        "CONFIDENTIAL — For internal use only. "
        "Generated automatically by ReqFlow AI Enterprise."
    )
    cfr.font.name = "Calibri"
    cfr.font.size = Pt(8.5)
    cfr.font.italic = True
    cfr.font.color.rgb = _CLR_DIM

    _add_page_break(doc)


def _add_metadata_table(
    doc: Document,
    project_name: str,
    industry: str,
    timestamp: str,
    section_label: str,
) -> None:
    """Renders a compact project metadata info table."""
    rows = [
        ("Project",   project_name),
        ("Industry",  industry),
        ("Document",  section_label),
        ("Generated", timestamp),
        ("Tool",      "ReqFlow AI — ERP Requirement Engine"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(rows):
        lc = table.cell(i, 0)
        vc = table.cell(i, 1)
        lc.width = Inches(1.8)
        vc.width = Inches(4.7)
        _set_cell_bg(lc, _CLR_TH_BG)
        _set_cell_border(lc)
        _set_cell_border(vc)
        lr = lc.paragraphs[0].add_run(label)
        lr.font.bold = True
        lr.font.size = Pt(9.5)
        lr.font.color.rgb = _CLR_TH_TEXT
        vr = vc.paragraphs[0].add_run(value)
        vr.font.size = Pt(9.5)
        vr.font.color.rgb = _CLR_BODY
    doc.add_paragraph()


def _add_section_header(doc: Document, title: str, code: str) -> None:
    """Adds a visually prominent section divider for the complete suite."""
    _add_horizontal_rule(doc)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    p = doc.add_heading(level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(f"[{code}]  {title}")
    r.font.color.rgb = _CLR_DARK
    r.font.size = Pt(18)
    r.font.bold = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(12)


def _doc_footer(doc: Document, project_name: str, label: str, timestamp: str) -> None:
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(
        f"— End of {label} —  "
        f"{project_name}  |  Generated by ReqFlow AI  |  {timestamp}"
    )
    r.font.size = Pt(8.5)
    r.font.italic = True
    r.font.color.rgb = _CLR_DIM
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _to_bytes(doc: Document) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ===========================================================================
#  Public API
# ===========================================================================

def build_brd_docx(
    project_name: str,
    industry: str,
    content: str,
    timestamp: Optional[str] = None,
) -> bytes:
    """
    Builds a standalone BRD Word document.

    Returns
    -------
    bytes   Ready for ``st.download_button(data=...)``
    """
    ts = timestamp or datetime.now().strftime("%B %d, %Y at %H:%M")
    doc = _create_base_doc()
    _add_cover_page(doc, project_name, industry, "Business Requirements Document (BRD)", ts)
    _add_metadata_table(doc, project_name, industry, ts, "Business Requirements Document (BRD)")
    _render_markdown(doc, content)
    _doc_footer(doc, project_name, "Business Requirements Document", ts)
    return _to_bytes(doc)


def build_section_docx(
    project_name: str,
    industry: str,
    section_key: str,
    content: str,
    timestamp: Optional[str] = None,
) -> bytes:
    """
    Builds a Word document for any single ERP section.

    Parameters
    ----------
    section_key : one of brd | srs | use_cases | user_stories |
                  db_suggestions | kpis | workflow | reports
    """
    ts = timestamp or datetime.now().strftime("%B %d, %Y at %H:%M")
    meta = _SECTION_META.get(
        section_key,
        (section_key.replace("_", " ").title(), section_key.upper()[:4])
    )
    title, code = meta
    doc = _create_base_doc()
    _add_cover_page(doc, project_name, industry, title, ts)
    _add_metadata_table(doc, project_name, industry, ts, title)
    _render_markdown(doc, content)
    _doc_footer(doc, project_name, title, ts)
    return _to_bytes(doc)


def build_complete_suite_docx(
    project_name: str,
    industry: str,
    timestamp: Optional[str],
    brd: str,
    srs: str,
    use_cases: str,
    user_stories: str,
    db_suggestions: str,
    kpis: str = "",
    workflow: str = "",
    reports: str = "",
) -> bytes:
    """
    Builds a single .docx containing all 8 ERP sections with:
    - Branded cover page
    - Table of contents listing
    - Metadata table
    - Each section separated by a page break
    - Section divider headings
    - Footer on final page

    Returns
    -------
    bytes   Ready for ``st.download_button(data=...)``
    """
    ts = timestamp or datetime.now().strftime("%B %d, %Y at %H:%M")
    doc = _create_base_doc()

    # Cover
    _add_cover_page(doc, project_name, industry, "Complete ERP Requirement Suite", ts)

    # Table of contents (plain list — auto TOC needs VBA, not supported in docx)
    p = doc.add_heading(level=1)
    r = p.add_run("Table of Contents")
    r.font.color.rgb = _CLR_DARK
    r.font.size = Pt(16)
    p.paragraph_format.space_after = Pt(10)

    all_sections = [
        ("brd",            brd),
        ("srs",            srs),
        ("use_cases",      use_cases),
        ("user_stories",   user_stories),
        ("db_suggestions", db_suggestions),
        ("kpis",           kpis),
        ("workflow",       workflow),
        ("reports",        reports),
    ]

    for i, (key, content) in enumerate(all_sections, 1):
        if not content or not content.strip():
            continue
        meta_title = _SECTION_META[key][0]
        toc_p = doc.add_paragraph()
        toc_r = toc_p.add_run(f"    {i}.  {meta_title}")
        toc_r.font.size = Pt(10.5)
        toc_r.font.color.rgb = _CLR_PURPLE
        toc_p.paragraph_format.space_after = Pt(3)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    _add_horizontal_rule(doc)
    _add_metadata_table(
        doc, project_name, industry, ts,
        "Complete ERP Requirement Suite (8 Sections)"
    )
    _add_page_break(doc)

    # Render sections
    included = [(k, c) for k, c in all_sections if c and c.strip()]
    for idx, (key, content) in enumerate(included):
        title, code = _SECTION_META[key]
        _add_section_header(doc, title, code)
        _render_markdown(doc, content)
        if idx < len(included) - 1:
            _add_page_break(doc)

    _doc_footer(doc, project_name, "Complete ERP Requirement Suite", ts)
    return _to_bytes(doc)
